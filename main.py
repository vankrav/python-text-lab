from nltk.tokenize import word_tokenize
from nltk.stem import PorterStemmer
from nltk.corpus import stopwords
from clean_text import clean_stopwords_punctuation
from tabulate import tabulate
from nltk.stem.snowball import SnowballStemmer
import docx
from pymystem3 import Mystem
m = Mystem()



# Чтение текста из файлов .docx
doc1 = docx.Document('1.docx')
text1 = ""
for paragraph in doc1.paragraphs:
    text1 += paragraph.text + "\n"

doc2 = docx.Document('2.docx')
text2 = ""
for paragraph in doc2.paragraphs:
    text2 += paragraph.text + "\n"

doc3 = docx.Document('3.docx')
text3 = ""
for paragraph in doc3.paragraphs:
    text3 += paragraph.text + "\n"

doc4 = docx.Document('test.docx')
text4 = ""
for paragraph in doc4.paragraphs:
    text4 += paragraph.text + "\n"



# Объединение первых трех текстов для создания корпуса
corpus = clean_stopwords_punctuation(text1 + text2 + text3)
# Подготовка четвертого текста
text_4 = clean_stopwords_punctuation(text4)

stemmer = SnowballStemmer("russian")

stemmed_corpus = [stemmer.stem(word) for word in corpus]
stemmed_text_4 = [stemmer.stem(word) for word in text_4]

# Подсчет частоты встречаемости слов в корпусе и тексте 4
word_frequency_corpus = {}
total_words_corpus = len(stemmed_corpus)

word_frequency_text_4 = {}
total_words_text_4 = len(stemmed_text_4)

for word in stemmed_corpus:
    word_frequency_corpus[word] = word_frequency_corpus.get(word, 0) + 1

for word in stemmed_text_4:
    word_frequency_text_4[word] = word_frequency_text_4.get(word, 0) + 1

sorted_words_corpus = sorted(word_frequency_corpus.items(), key=lambda x: x[1], reverse=True)

# Подсчет количества общих слов
common_words = set(word_frequency_text_4.keys()) & set(word_frequency_corpus.keys())
total_common_words = sum(word_frequency_text_4[word] for word in common_words)

# Вычисление процента схожести слов
percentage_common_words = total_common_words / total_words_text_4 * 100

# порог 80 по умолчанию
percentage = 80

if percentage_common_words > percentage:
    print("Более {}% слов из четвертого текста встречаются в корпусе, тексты схожей тематики.".format(percentage))
else:
    print("Менее {}% слов из четвертого текста встречаются в корпусе, тексты различной тематики.".format(percentage))
# количество строк
number_of_rows = 15
k = 0
table_data = []

for word, frequency in sorted_words_corpus:
    if k >= number_of_rows:
        break
    k+=1

    percent_in_corpus = frequency / total_words_corpus * 100
    percent_in_text_4 = word_frequency_text_4.get(word, 0) / total_words_text_4 * 100
    word = ''.join(m.lemmatize(word))
    table_data.append([word, f"{percent_in_corpus:.2f}%", f"{percent_in_text_4:.2f}%"])

# Вывод таблицы
headers = ["Термин", "% в корпусе", "% в 4 тексте"]
print(tabulate(table_data, headers=headers, tablefmt="pretty"))



